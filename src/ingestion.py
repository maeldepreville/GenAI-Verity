import logging
import re
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Optional, BinaryIO

import chardet
import pdfplumber
from docx import Document
from pypdf import PdfReader


logger = logging.getLogger(__name__)


class DocumentType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    UNKNOWN = "unknown"


class ExtractionStatus(Enum):
    SUCCESS = "success"
    PARTIAL = "partial"
    FAILED = "failed"


@dataclass
class DocumentMetadata:
    filename: str
    file_type: DocumentType
    file_size: int
    page_count: Optional[int]
    word_count: int
    char_count: int
    author: Optional[str]
    created_date: Optional[datetime]
    encoding: Optional[str]


@dataclass
class ExtractionResult:
    text: str
    metadata: DocumentMetadata
    status: ExtractionStatus
    error_message: Optional[str]
    extraction_method: str


class DocumentParser:
    MAX_FILE_SIZE = 50 * 1024 * 1024
    SUPPORTED_TYPES = {DocumentType.PDF, DocumentType.DOCX, DocumentType.TXT}

    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)

    def parse(self, file: BinaryIO, filename: str) -> ExtractionResult:
        try:
            self._validate_file_size(file)
            doc_type = self._detect_document_type(filename)
            self._validate_document_type(doc_type)

            if doc_type == DocumentType.PDF:
                return self._parse_pdf(file, filename)
            elif doc_type == DocumentType.DOCX:
                return self._parse_docx(file, filename)
            elif doc_type == DocumentType.TXT:
                return self._parse_txt(file, filename)

        except Exception as e:
            self.logger.error(f"Failed to parse {filename}: {str(e)}")
            return self._create_error_result(filename, str(e))

    def _validate_file_size(self, file: BinaryIO) -> None:
        file.seek(0, 2)
        size = file.tell()
        file.seek(0)

        if size > self.MAX_FILE_SIZE:
            raise ValueError(f"File size {size} exceeds maximum {self.MAX_FILE_SIZE}")

    def _detect_document_type(self, filename: str) -> DocumentType:
        extension = Path(filename).suffix.lower().lstrip('.')

        try:
            return DocumentType(extension)
        except ValueError:
            return DocumentType.UNKNOWN

    def _validate_document_type(self, doc_type: DocumentType) -> None:
        if doc_type not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported document type: {doc_type.value}")

    def _parse_pdf(self, file: BinaryIO, filename: str) -> ExtractionResult:
        try:
            text, metadata = self._extract_pdf_pdfplumber(file, filename)
            method = "pdfplumber"
        except Exception as e:
            self.logger.warning(f"pdfplumber failed, trying pypdf: {str(e)}")
            file.seek(0)
            text, metadata = self._extract_pdf_pypdf(file, filename)
            method = "pypdf"

        preprocessed_text = self._preprocess_text(text)

        metadata.word_count = self._count_words(preprocessed_text)
        metadata.char_count = len(preprocessed_text)

        return ExtractionResult(
            text=preprocessed_text,
            metadata=metadata,
            status=ExtractionStatus.SUCCESS,
            error_message=None,
            extraction_method=method
        )

    def _extract_pdf_pdfplumber(self, file: BinaryIO, filename: str) -> tuple[str, DocumentMetadata]:
        file.seek(0)
        text_parts = []

        with pdfplumber.open(file) as pdf:
            page_count = len(pdf.pages)

            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            metadata_dict = pdf.metadata or {}
            author = metadata_dict.get('Author')
            created = metadata_dict.get('CreationDate')

            created_date = None
            if created:
                try:
                    created_date = self._parse_pdf_date(created)
                except Exception:
                    pass

        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        metadata = DocumentMetadata(
            filename=filename,
            file_type=DocumentType.PDF,
            file_size=file_size,
            page_count=page_count,
            word_count=0,
            char_count=0,
            author=author,
            created_date=created_date,
            encoding=None
        )

        return "\n".join(text_parts), metadata

    def _extract_pdf_pypdf(self, file: BinaryIO, filename: str) -> tuple[str, DocumentMetadata]:
        file.seek(0)
        reader = PdfReader(file)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        metadata_dict = reader.metadata or {}
        author = metadata_dict.get('/Author')
        created = metadata_dict.get('/CreationDate')

        created_date = None
        if created:
            try:
                created_date = self._parse_pdf_date(created)
            except Exception:
                pass

        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        metadata = DocumentMetadata(
            filename=filename,
            file_type=DocumentType.PDF,
            file_size=file_size,
            page_count=len(reader.pages),
            word_count=0,
            char_count=0,
            author=author,
            created_date=created_date,
            encoding=None
        )

        return "\n".join(text_parts), metadata

    def _parse_docx(self, file: BinaryIO, filename: str) -> ExtractionResult:
        file.seek(0)
        doc = Document(file)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        text = "\n".join(text_parts)
        preprocessed_text = self._preprocess_text(text)

        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        core_props = doc.core_properties
        author = core_props.author if core_props.author else None
        created = core_props.created if core_props.created else None

        metadata = DocumentMetadata(
            filename=filename,
            file_type=DocumentType.DOCX,
            file_size=file_size,
            page_count=None,
            word_count=self._count_words(preprocessed_text),
            char_count=len(preprocessed_text),
            author=author,
            created_date=created,
            encoding=None
        )

        return ExtractionResult(
            text=preprocessed_text,
            metadata=metadata,
            status=ExtractionStatus.SUCCESS,
            error_message=None,
            extraction_method="python-docx"
        )

    def _parse_txt(self, file: BinaryIO, filename: str) -> ExtractionResult:
        file.seek(0)
        raw_data = file.read()

        detection = chardet.detect(raw_data)
        encoding = detection['encoding'] or 'utf-8'

        try:
            text = raw_data.decode(encoding)
        except UnicodeDecodeError:
            text = raw_data.decode('utf-8', errors='ignore')
            encoding = 'utf-8-fallback'

        preprocessed_text = self._preprocess_text(text)

        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        metadata = DocumentMetadata(
            filename=filename,
            file_type=DocumentType.TXT,
            file_size=file_size,
            page_count=None,
            word_count=self._count_words(preprocessed_text),
            char_count=len(preprocessed_text),
            author=None,
            created_date=None,
            encoding=encoding
        )

        return ExtractionResult(
            text=preprocessed_text,
            metadata=metadata,
            status=ExtractionStatus.SUCCESS,
            error_message=None,
            extraction_method="chardet"
        )

    def _preprocess_text(self, text: str) -> str:
        text = self._normalize_whitespace(text)
        text = self._remove_page_numbers(text)
        text = self._fix_hyphenation(text)
        text = self._normalize_unicode(text)
        return text.strip()

    def _normalize_whitespace(self, text: str) -> str:
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text

    def _remove_page_numbers(self, text: str) -> str:
        text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
        return text

    def _fix_hyphenation(self, text: str) -> str:
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        return text

    def _normalize_unicode(self, text: str) -> str:
        replacements = {
            '\u2018': "'",
            '\u2019': "'",
            '\u201c': '"',
            '\u201d': '"',
            '\u2013': '-',
            '\u2014': '--',
            '\u2026': '...',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    def _count_words(self, text: str) -> int:
        words = re.findall(r'\b\w+\b', text)
        return len(words)

    def _parse_pdf_date(self, date_str: str) -> datetime:
        if date_str.startswith('D:'):
            date_str = date_str[2:]

        date_str = date_str[:14]

        return datetime.strptime(date_str, '%Y%m%d%H%M%S')

    def _create_error_result(self, filename: str, error_msg: str) -> ExtractionResult:
        metadata = DocumentMetadata(
            filename=filename,
            file_type=DocumentType.UNKNOWN,
            file_size=0,
            page_count=None,
            word_count=0,
            char_count=0,
            author=None,
            created_date=None,
            encoding=None
        )

        return ExtractionResult(
            text="",
            metadata=metadata,
            status=ExtractionStatus.FAILED,
            error_message=error_msg,
            extraction_method="none"
        )


class TextSegmenter:
    def __init__(self, max_tokens: int = 4000, overlap: int = 200):
        self.max_tokens = max_tokens
        self.overlap = overlap

    def segment(self, text: str) -> list[str]:
        paragraphs = self._split_into_paragraphs(text)
        segments = self._create_segments(paragraphs)
        return segments

    def _split_into_paragraphs(self, text: str) -> list[str]:
        paragraphs = text.split('\n\n')
        return [p.strip() for p in paragraphs if p.strip()]

    def _create_segments(self, paragraphs: list[str]) -> list[str]:
        segments = []
        current_segment = []
        current_length = 0

        for para in paragraphs:
            para_length = len(para)

            if current_length + para_length > self.max_tokens and current_segment:
                segments.append('\n\n'.join(current_segment))

                overlap_text = self._get_overlap_text(current_segment)
                current_segment = [overlap_text] if overlap_text else []
                current_length = len(overlap_text) if overlap_text else 0

            current_segment.append(para)
            current_length += para_length

        if current_segment:
            segments.append('\n\n'.join(current_segment))

        return segments

    def _get_overlap_text(self, segments: list[str]) -> str:
        combined = '\n\n'.join(segments)
        if len(combined) <= self.overlap:
            return combined

        return combined[-self.overlap:]


